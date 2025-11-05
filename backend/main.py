from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import uvicorn
import time
import uuid
from typing import Dict, List, Optional
from pydantic import BaseModel
from google import genai
import asyncio
import os
from pathlib import Path
import PyPDF2
import io
try:
    from docx import Document
except ImportError:
    Document = None

# Initialize FastAPI app
app = FastAPI(title="LearnLens API", version="1.0.0")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Gemini client
client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

# In-memory storage (replace with database in production)
documents: Dict[str, dict] = {}
videos: Dict[str, dict] = {}
users: Dict[str, dict] = {"default": {"isPremium": False, "freeVideosUsed": 0}}

class Topic(BaseModel):
    id: str
    title: str
    description: str
    order: int
    isPremium: bool

class DocumentOutline(BaseModel):
    id: str
    title: str
    topics: List[Topic]
    createdAt: str

class VideoGenerationRequest(BaseModel):
    topicId: str

def extract_text_from_file(content: bytes, content_type: str, filename: str) -> str:
    """Extract text from different file types"""
    print(f"\n=== CONTENT EXTRACTION START ===")
    print(f"File: {filename}")
    print(f"Content Type: {content_type}")
    print(f"File Size: {len(content)} bytes")
    
    try:
        extracted_text = ""
        
        if content_type == "text/plain":
            print("Processing as plain text file...")
            extracted_text = content.decode('utf-8')
        
        elif content_type == "application/pdf":
            print("Processing as PDF file...")
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            print(f"PDF has {len(pdf_reader.pages)} pages")
            
            for i, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                extracted_text += page_text + "\n"
                print(f"Page {i+1}: Extracted {len(page_text)} characters")
        
        elif content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            print("Processing as DOCX file...")
            if Document is None:
                print("WARNING: python-docx not installed, cannot extract DOCX content")
                return "DOCX file uploaded - content extraction not available"
            
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)
            print(f"DOCX has {len(doc.paragraphs)} paragraphs")
            
            for i, paragraph in enumerate(doc.paragraphs):
                extracted_text += paragraph.text + "\n"
                if i < 3:  # Log first 3 paragraphs
                    print(f"Paragraph {i+1}: {paragraph.text[:100]}...")
        
        elif content_type == "application/msword":
            print("Processing as DOC file (basic support)...")
            extracted_text = content.decode('utf-8', errors='ignore')
        
        else:
            raise ValueError(f"Unsupported file type: {content_type}")
        
        print(f"\n--- EXTRACTED CONTENT PREVIEW ---")
        print(f"Total characters extracted: {len(extracted_text)}")
        print(f"First 500 characters:")
        print(f"{extracted_text[:500]}...")
        print(f"=== CONTENT EXTRACTION END ===\n")
        
        return extracted_text
            
    except Exception as e:
        print(f"ERROR extracting text from {filename}: {e}")
        print(f"=== CONTENT EXTRACTION FAILED ===\n")
        return f"Error reading file: {filename}"

@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload and process document"""
    print(f"\n🚀 === DOCUMENT UPLOAD STARTED ===")
    print(f"File: {file.filename}")
    print(f"Content Type: {file.content_type}")
    print(f"Timestamp: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    
    try:
        # Validate file type
        allowed_types = ["application/pdf", "text/plain", "application/msword", 
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
        
        print(f"\n📋 Step 1: File validation")
        print(f"Allowed types: {allowed_types}")
        print(f"File type: {file.content_type}")
        
        if file.content_type not in allowed_types:
            print(f"❌ File type not supported: {file.content_type}")
            raise HTTPException(status_code=400, detail="Unsupported file type")
        
        print(f"✅ File type validation passed")
        
        # Read file content
        print(f"\n📖 Step 2: Reading file content")
        content = await file.read()
        print(f"File size: {len(content)} bytes")
        
        # Generate document ID
        doc_id = str(uuid.uuid4())
        print(f"Generated document ID: {doc_id}")
        
        # Extract text content
        print(f"\n🔍 Step 3: Extracting text content")
        text_content = extract_text_from_file(content, file.content_type, file.filename)
        
        if not text_content.strip():
            print(f"❌ No text content extracted from file")
            raise HTTPException(status_code=400, detail="Could not extract text from file")
        
        print(f"✅ Text extraction completed: {len(text_content)} characters")
        
        # Generate outline using Gemini
        print(f"\n🤖 Step 4: Generating outline with AI")
        outline = await generate_outline(text_content, file.filename)
        print(f"✅ Outline generated with {len(outline['topics'])} topics")
        
        # Store document
        print(f"\n💾 Step 5: Storing document")
        documents[doc_id] = {
            "id": doc_id,
            "filename": file.filename,
            "content": text_content,
            "outline": outline,
            "createdAt": time.time()
        }
        print(f"✅ Document stored successfully")
        
        # Generate first video (free)
        print(f"\n🎬 Step 6: Generating first video (free)")
        first_topic = outline["topics"][0]
        print(f"First topic: {first_topic['title']}")
        video_id = await generate_video_async(first_topic["id"], first_topic["title"], text_content)
        print(f"✅ Video generation started with ID: {video_id}")
        
        print(f"\n🎉 === DOCUMENT UPLOAD COMPLETED ===")
        print(f"Document ID: {doc_id}")
        print(f"Topics generated: {len(outline['topics'])}")
        print(f"First video ID: {video_id}")
        
        return JSONResponse({
            "documentId": doc_id,
            "outline": outline
        })
        
    except Exception as e:
        print(f"\n❌ === DOCUMENT UPLOAD FAILED ===")
        print(f"Error: {str(e)}")
        print(f"Error type: {type(e).__name__}")
        raise HTTPException(status_code=500, detail=str(e))

async def generate_outline(content: str, filename: str) -> dict:
    """Generate document outline using Gemini"""
    print(f"\n=== OUTLINE GENERATION START ===")
    print(f"File: {filename}")
    print(f"Content length: {len(content)} characters")
    
    # Check if API key is set
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        print("ERROR: GEMINI_API_KEY not set in environment variables")
        print("=== OUTLINE GENERATION FAILED ===\n")
        return generate_fallback_outline(filename)
    
    try:
        # Prepare content for LLM (limit to 2000 chars for prompt)
        content_preview = content[:2000]
        print(f"\n--- CONTENT SENT TO LLM ---")
        print(f"Content preview (first 2000 chars):")
        print(f"{content_preview}...")
        
        prompt = f"""
        Analyze the following document content and create a structured learning outline.
        
        Document: {filename}
        Content: {content_preview}...
        
        Create 5-7 main topics that would make good educational videos. For each topic, provide:
        1. A clear, engaging title
        2. A brief description of what the video would cover
        3. Logical order for learning
        
        Format as JSON with this structure:
        {{
            "title": "Document Title",
            "topics": [
                {{
                    "title": "Topic Title",
                    "description": "What this video covers"
                }}
            ]
        }}
        """
        
        print(f"\n--- SENDING TO GEMINI ---")
        print(f"Model: gemini-2.5-flash")
        print(f"Prompt length: {len(prompt)} characters")
        
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt
        )
        
        outline_id = str(uuid.uuid4())
        topics = []
        
        # Use Gemini response to generate topics based on actual content
        gemini_text = response.text
        print(f"\n--- GEMINI RESPONSE ---")
        print(f"Response length: {len(gemini_text)} characters")
        print(f"Full Gemini response:")
        print(f"{gemini_text}")
        print(f"--- END GEMINI RESPONSE ---\n")
        
        # Extract topics from Gemini response (improved parsing)
        lines = gemini_text.split('\n')
        topic_count = 0
        
        for line in lines:
            if topic_count >= 5:  # Limit to 5 topics
                break
            
            # Look for numbered items, bullet points, or titles
            line = line.strip()
            if not line or len(line) < 5:
                continue
                
            # Try different patterns
            title = None
            if line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.')):
                title = line.split('.', 1)[1].strip()
            elif line.startswith(('-', '*', '•')):
                title = line[1:].strip()
            elif ':' in line and len(line.split(':', 1)) == 2:
                title = line.split(':', 1)[1].strip()
            elif line.startswith('##') or line.startswith('**'):
                title = line.replace('#', '').replace('*', '').strip()
            
            if title and len(title) > 3:
                description = f"Educational content about {title.lower()}"
                
                topic_id = str(uuid.uuid4())
                topics.append({
                    "id": topic_id,
                    "title": title,
                    "description": description,
                    "order": topic_count,
                    "isPremium": topic_count > 0,  # First video is free
                    "video": None
                })
                topic_count += 1
                print(f"Extracted topic {topic_count}: {title}")
        
        print(f"Total topics extracted: {len(topics)}")
        
        # Fallback if no topics extracted
        if not topics:
            print("No topics extracted from Gemini, using fallback")
            for i, topic in enumerate([
                {"title": "Introduction and Overview", "description": "Get started with the fundamentals"},
                {"title": "Core Concepts", "description": "Understanding the main principles"},
                {"title": "Practical Applications", "description": "Real-world examples and use cases"}
            ]):
                topic_id = str(uuid.uuid4())
                topics.append({
                    "id": topic_id,
                    "title": topic["title"],
                    "description": topic["description"],
                    "order": i,
                    "isPremium": i > 0,  # First video is free
                    "video": None
                })
        
        return {
            "id": outline_id,
            "title": filename.replace('.pdf', '').replace('.txt', '').replace('.doc', ''),
            "topics": topics,
            "createdAt": time.strftime("%Y-%m-%d %H:%M:%S")
        }
        
    except Exception as e:
        print(f"Error generating outline: {e}")
        return generate_fallback_outline(filename)

def generate_fallback_outline(filename: str) -> dict:
    """Generate fallback outline when Gemini fails"""
    print("Using fallback outline generation")
    return {
        "id": str(uuid.uuid4()),
        "title": filename.replace('.pdf', '').replace('.txt', '').replace('.doc', ''),
        "topics": [{
            "id": str(uuid.uuid4()),
            "title": "Introduction",
            "description": "Overview of the document content",
            "order": 0,
            "isPremium": False,
            "video": None
        }],
        "createdAt": time.strftime("%Y-%m-%d %H:%M:%S")
    }

async def generate_video_async(topic_id: str, title: str, content: str) -> str:
    """Generate video using Veo (simulated)"""
    video_id = str(uuid.uuid4())
    
    print(f"\n🎬 Starting async video generation")
    print(f"Video ID: {video_id}")
    print(f"Topic ID: {topic_id}")
    print(f"Title: {title}")
    
    # Store video with generating status
    videos[video_id] = {
        "id": video_id,
        "topicId": topic_id,
        "status": "generating",
        "title": title,
        "createdAt": time.time()
    }
    
    print(f"✅ Video record created with 'generating' status")
    
    # Simulate async video generation
    print(f"🚀 Starting background video generation task")
    asyncio.create_task(simulate_video_generation(video_id, title, content))
    
    return video_id

async def generate_video_script(title: str, content: str) -> str:
    """Generate video script using Gemini"""
    print(f"\n=== VIDEO SCRIPT GENERATION START ===")
    print(f"Topic: {title}")
    print(f"Content length: {len(content)} characters")
    
    try:
        content_preview = content[:1000]
        print(f"\n--- CONTENT FOR SCRIPT ---")
        print(f"Content preview (first 1000 chars):")
        print(f"{content_preview}...")
        
        prompt = f"""
        Create an engaging educational video script for the topic: "{title}"
        
        Based on this content: {content_preview}...
        
        The script should:
        1. Be 2-3 minutes long when spoken
        2. Include clear explanations and examples
        3. Be engaging and educational
        4. Have a clear introduction, main content, and conclusion
        
        Format as a natural speaking script.
        """
        
        print(f"\n--- SENDING SCRIPT REQUEST TO GEMINI ---")
        print(f"Prompt length: {len(prompt)} characters")
        
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt
        )
        
        script_text = response.text
        print(f"\n--- SCRIPT RESPONSE ---")
        print(f"Script length: {len(script_text)} characters")
        print(f"Generated script:")
        print(f"{script_text}")
        print(f"--- END SCRIPT RESPONSE ---")
        print(f"=== VIDEO SCRIPT GENERATION END ===\n")
        
        return script_text
        
    except Exception as e:
        print(f"ERROR generating script: {e}")
        print(f"=== VIDEO SCRIPT GENERATION FAILED ===\n")
        return f"Educational content about {title}"

async def simulate_video_generation(video_id: str, title: str, content: str):
    """Generate video using Gemini script + Veo"""
    print(f"\n=== VIDEO GENERATION START ===")
    print(f"Video ID: {video_id}")
    print(f"Topic: {title}")
    
    try:
        # Step 1: Generate script with Gemini
        print(f"Step 1: Generating script with Gemini...")
        script = await generate_video_script(title, content)
        
        # Step 2: Simulate Veo video generation
        print(f"\nStep 2: Simulating Veo video generation...")
        print(f"Waiting 5 seconds for demo video generation...")
        await asyncio.sleep(5)  # 5 seconds for demo
        
        # In production, this would call Veo API with the script:
        veo_prompt = f"Create an educational video with this script: {script[:500]}..."
        print(f"\n--- VEO PROMPT (SIMULATED) ---")
        print(f"Prompt for Veo: {veo_prompt}")
        print(f"--- END VEO PROMPT ---")
        
        # operation = client.models.generate_videos(model="veo-2.0-generate-001", prompt=veo_prompt)
        
        # Update video status
        videos[video_id].update({
            "status": "ready",
            "url": "https://commondatastorage.googleapis.com/gtv-videos-bucket/sample/BigBuckBunny.mp4",  # Demo video
            "duration": 180,  # 3 minutes
            "thumbnail": "https://commondatastorage.googleapis.com/gtv-videos-bucket/sample/images/BigBuckBunny.jpg",
            "script": script
        })
        
        print(f"\nVideo generation completed successfully!")
        print(f"Video URL: {videos[video_id]['url']}")
        print(f"=== VIDEO GENERATION END ===\n")
        
    except Exception as e:
        videos[video_id]["status"] = "error"
        print(f"ERROR: Video generation failed: {e}")
        print(f"=== VIDEO GENERATION FAILED ===\n")

@app.get("/documents/{document_id}")
async def get_document_status(document_id: str):
    """Get document outline and video status"""
    if document_id not in documents:
        raise HTTPException(status_code=404, detail="Document not found")
    
    doc = documents[document_id]
    outline = doc["outline"].copy()
    
    # Update video status for each topic
    for topic in outline["topics"]:
        topic_videos = [v for v in videos.values() if v["topicId"] == topic["id"]]
        if topic_videos:
            video = topic_videos[0]
            topic["video"] = {
                "id": video["id"],
                "status": video["status"],
                "url": video.get("url"),
                "duration": video.get("duration"),
                "thumbnail": video.get("thumbnail")
            }
    
    return outline

@app.post("/generate-video")
async def generate_video(request: VideoGenerationRequest):
    """Generate video for a topic"""
    # Find the topic
    topic_found = False
    topic_title = ""
    content = ""
    
    for doc in documents.values():
        for topic in doc["outline"]["topics"]:
            if topic["id"] == request.topicId:
                topic_found = True
                topic_title = topic["title"]
                content = doc["content"]
                break
        if topic_found:
            break
    
    if not topic_found:
        raise HTTPException(status_code=404, detail="Topic not found")
    
    # Check if user can generate video (premium check)
    user = users["default"]
    if not user["isPremium"] and user["freeVideosUsed"] >= 1:
        raise HTTPException(status_code=403, detail="Premium subscription required")
    
    # Generate video
    video_id = await generate_video_async(request.topicId, topic_title, content)
    
    # Update free videos used
    if not user["isPremium"]:
        user["freeVideosUsed"] += 1
    
    return {"videoId": video_id}

@app.get("/videos/{video_id}")
async def get_video_status(video_id: str):
    """Get video generation status"""
    if video_id not in videos:
        raise HTTPException(status_code=404, detail="Video not found")
    
    return videos[video_id]

@app.get("/subscription")
async def check_subscription():
    """Check user subscription status"""
    user = users["default"]
    return {
        "isPremium": user["isPremium"],
        "freeVideosUsed": user["freeVideosUsed"]
    }

@app.post("/subscribe")
async def subscribe(plan: str):
    """Handle subscription (mock)"""
    user = users["default"]
    user["isPremium"] = True
    return {"success": True, "plan": plan}

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)